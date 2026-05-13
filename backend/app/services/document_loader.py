"""Filesystem-backed document loader for the Document Handler node.

Lists and reads documents from an analyst-supplied absolute path.
Supports `.txt`, `.md`, `.pdf` (via pypdf), and `.docx` (via
python-docx); each file is converted to plain text before being handed
to the executor. Intentionally simple: no recursion into subdirectories,
no sandboxing beyond extension filtering and a per-file size cap. The
dev tool is local-first with no auth — analysts pointing at their own
filesystem are trusted.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

_ALLOWED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}
_MAX_BYTES = 25 * 1024 * 1024  # 25MB — PDFs/docx are bigger than plain text


class DocumentLoaderError(Exception):
    """Raised when a path is invalid or a file can't be served."""


@dataclass(frozen=True)
class DocumentEntry:
    name: str
    size: int

    def as_dict(self) -> dict[str, int | str]:
        return {"name": self.name, "size": self.size}


def _resolve_dir(path: str) -> Path:
    if not path:
        raise DocumentLoaderError("path is required")
    p = Path(path)
    if not p.is_absolute():
        raise DocumentLoaderError("path must be absolute")
    if not p.exists():
        raise DocumentLoaderError(f"path does not exist: {path}")
    if not p.is_dir():
        raise DocumentLoaderError(f"path is not a directory: {path}")
    return p


def list_files(path: str) -> list[DocumentEntry]:
    """Top-level files matching allowed extensions. Sorted by name."""
    directory = _resolve_dir(path)
    entries: list[DocumentEntry] = []
    for child in sorted(directory.iterdir(), key=lambda p: p.name.lower()):
        if not child.is_file():
            continue
        if child.suffix.lower() not in _ALLOWED_EXTENSIONS:
            continue
        try:
            size = child.stat().st_size
        except OSError:
            continue
        entries.append(DocumentEntry(name=child.name, size=size))
    return entries


def read_file(path: str, filename: str) -> str:
    """Read a single file's text content. Capped at _MAX_BYTES on disk."""
    directory = _resolve_dir(path)
    if not filename or "/" in filename or "\\" in filename or filename.startswith("."):
        raise DocumentLoaderError("invalid filename")
    target = directory / filename
    if not target.is_file():
        raise DocumentLoaderError(f"file not found: {filename}")
    suffix = target.suffix.lower()
    if suffix not in _ALLOWED_EXTENSIONS:
        raise DocumentLoaderError(f"unsupported file type: {target.suffix}")
    if target.stat().st_size > _MAX_BYTES:
        raise DocumentLoaderError(
            f"file too large (>{_MAX_BYTES // (1024 * 1024)} MB): {filename}"
        )
    if suffix == ".pdf":
        return _read_pdf(target)
    if suffix == ".docx":
        return _read_docx(target)
    return target.read_text(encoding="utf-8", errors="replace")


def _read_pdf(target: Path) -> str:
    """Extract page text from a PDF. Empty pages become blank lines.

    pypdf is a pure-Python reader — fine for the text-extraction-only
    case the regulatory workflow needs. Scanned/image PDFs without an
    embedded text layer will return very little; OCR is out of scope.
    """
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError as exc:
        raise DocumentLoaderError(
            "PDF support requires pypdf; install backend/requirements.txt"
        ) from exc
    try:
        reader = PdfReader(str(target))
        if reader.is_encrypted:
            raise DocumentLoaderError(f"encrypted PDF not supported: {target.name}")
        pages = [page.extract_text() or "" for page in reader.pages]
    except PdfReadError as exc:
        raise DocumentLoaderError(f"could not read PDF {target.name}: {exc}") from exc
    return "\n\n".join(p.strip() for p in pages).strip()


def _read_docx(target: Path) -> str:
    """Extract paragraph text from a Word .docx file.

    .doc (legacy binary) is not supported — python-docx only handles the
    XML-based .docx format. Tables and headers are ignored for now.
    """
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError as exc:
        raise DocumentLoaderError(
            "Word support requires python-docx; install backend/requirements.txt"
        ) from exc
    try:
        doc = Document(str(target))
    except PackageNotFoundError as exc:
        raise DocumentLoaderError(
            f"could not read Word file {target.name}: not a valid .docx"
        ) from exc
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return "\n\n".join(paragraphs)


@dataclass(frozen=True)
class LoadedDocument:
    filename: str
    text: str


def read_all_files(path: str) -> list[LoadedDocument]:
    """Read every supported file in the directory in name order."""
    entries = list_files(path)
    out: list[LoadedDocument] = []
    for entry in entries:
        out.append(LoadedDocument(filename=entry.name, text=read_file(path, entry.name)))
    return out
